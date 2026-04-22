"""
RightCut — DOCX parser using python-docx.
"""

from __future__ import annotations

import io
import logging

from models import ParsedDocument

logger = logging.getLogger(__name__)

MAX_CONTENT_CHARS = 50_000


async def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text and tables from a DOCX file."""
    try:
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(file_bytes))
        content_parts: list[str] = []
        tables: list[list[list[str]]] = []

        # Extract paragraphs (skip empty)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style.name.startswith("Heading"):
                    content_parts.append(f"\n## {text}\n")
                else:
                    content_parts.append(text)

        # Extract tables (one level deep)
        for tbl in doc.tables:
            table_data: list[list[str]] = []
            for row in tbl.rows:
                row_data = []
                seen_cells: set[int] = set()  # track merged cells by id
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue  # skip merged cells
                    seen_cells.add(cell_id)
                    row_data.append(cell.text.strip())
                if any(row_data):
                    table_data.append(row_data)

            if len(table_data) >= 2:
                tables.append(table_data)
                # Also include table content in the text
                content_parts.append("\n[Table]\n" + _table_to_text(table_data))

        # Extract headers/footers
        for section in doc.sections:
            try:
                header_text = "\n".join(
                    p.text.strip() for p in section.header.paragraphs if p.text.strip()
                )
                if header_text:
                    content_parts.insert(0, f"[Document Header]\n{header_text}\n")
            except Exception:
                pass

        content = "\n".join(content_parts)[:MAX_CONTENT_CHARS]

        return ParsedDocument(
            filename=filename,
            content=content,
            tables=tables,
            page_count=None,  # DOCX doesn't expose page count easily
            file_type="docx",
        )

    except Exception as e:
        logger.error(f"DOCX parsing failed for {filename}: {e}")
        return ParsedDocument(
            filename=filename,
            content=f"[Error parsing DOCX: {e}]",
            tables=[],
            page_count=None,
            file_type="docx",
        )


def _table_to_text(table: list[list[str]]) -> str:
    lines = []
    for row in table:
        lines.append(" | ".join(row))
    return "\n".join(lines)
